"""IMRD 学术报告自动生成模块。"""

from __future__ import annotations

import json
import logging
import re
import time
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from src.core.module_base import BaseModule
from src.research.phase_result import get_phase_value

logger = logging.getLogger(__name__)


_ANALYZE_METHOD_LABELS = {
    "reasoning_engine": "规则/证据驱动推理",
    "statistical_data_miner": "统计数据挖掘",
    "frequency_chi_square": "频次与卡方检验",
    "association_rules": "关联规则分析",
}


class ReportFormat(str, Enum):
    """支持的报告输出格式。"""

    MARKDOWN = "markdown"
    DOCX = "docx"


@dataclass
class Report:
    """结构化报告对象。"""

    title: str
    format: str
    content: str
    sections: Dict[str, str] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)
    output_path: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return {
            "title": self.title,
            "format": self.format,
            "content": self.content,
            "sections": dict(self.sections),
            "metadata": dict(self.metadata),
            "output_path": self.output_path,
        }


class ReportGenerator(BaseModule):
    """基于科研 session 结果生成 IMRD 学术报告。"""

    SECTION_ORDER = ("introduction", "methods", "results", "discussion")

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__("report_generator", config)
        self.output_dir = Path((config or {}).get("output_dir", "./output"))
        self.min_markdown_length = int((config or {}).get("min_markdown_length", 500))

    def _do_initialize(self) -> bool:
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.logger.info("报告生成器初始化完成")
        return True

    def _do_execute(self, context: Dict[str, Any]) -> Dict[str, Any]:
        session_result = context.get("session_result") or {}
        report_format = context.get("format", ReportFormat.MARKDOWN.value)
        report = self.generate_report(session_result, report_format)
        return {"report": report.to_dict()}

    def _do_cleanup(self) -> bool:
        self.logger.info("报告生成器清理完成")
        return True

    def generate_report(
        self,
        session_result: Dict[str, Any],
        report_format: str | ReportFormat,
    ) -> Report:
        """根据科研 session 结果生成报告。"""
        report_format = self._resolve_format(report_format)
        title = self._resolve_title(session_result)
        sections = {
            "introduction": self._build_introduction(session_result),
            "methods": self._build_methods(session_result),
            "results": self._build_results(session_result),
            "discussion": self._build_discussion(session_result),
        }
        markdown_content = self._render_markdown(title, session_result, sections)
        markdown_content = self._ensure_minimum_length(markdown_content, session_result, sections)

        output_path = ""
        if report_format is ReportFormat.MARKDOWN:
            output_path = self._render_markdown_file(session_result, markdown_content)
        elif report_format is ReportFormat.DOCX:
            output_path = self._render_docx(title, session_result, sections, markdown_content)

        metadata = {
            "generated_at": datetime.now().isoformat(),
            "session_id": self._resolve_string(session_result.get("session_id")) or f"session_{int(time.time())}",
            "research_question": self._resolve_question(session_result),
            "char_count": len(markdown_content),
            "section_count": len(sections),
            "output_format": report_format.value,
        }

        return Report(
            title=title,
            format=report_format.value,
            content=markdown_content,
            sections=sections,
            metadata=metadata,
            output_path=output_path,
        )

    def _resolve_format(self, value: str | ReportFormat) -> ReportFormat:
        if isinstance(value, ReportFormat):
            return value
        normalized = str(value).strip().lower()
        for item in ReportFormat:
            if item.value == normalized:
                return item
        raise ValueError(f"不支持的报告格式: {value!r}，可选: {[item.value for item in ReportFormat]}")

    def _resolve_title(self, session_result: Dict[str, Any]) -> str:
        metadata = session_result.get("metadata", {}) if isinstance(session_result.get("metadata"), dict) else {}
        question = self._resolve_question(session_result)
        return (
            self._resolve_string(metadata.get("title"))
            or self._resolve_string(session_result.get("title"))
            or f"中医科研 IMRD 报告：{question}"
        )

    def _resolve_question(self, session_result: Dict[str, Any]) -> str:
        metadata = session_result.get("metadata", {}) if isinstance(session_result.get("metadata"), dict) else {}
        return (
            self._resolve_string(session_result.get("question"))
            or self._resolve_string(session_result.get("research_question"))
            or self._resolve_string(metadata.get("research_question"))
            or "中医科研问题待补充"
        )

    def _phase_result(self, session_result: Dict[str, Any], phase_name: str) -> Dict[str, Any]:
        phase_results = session_result.get("phase_results")
        if isinstance(phase_results, dict):
            candidate = phase_results.get(phase_name)
            if isinstance(candidate, dict):
                return candidate
        for key in (phase_name, f"{phase_name}_phase", f"{phase_name}_result", f"{phase_name}_results"):
            candidate = session_result.get(key)
            if isinstance(candidate, dict):
                return candidate
        return {}

    def _build_introduction(self, session_result: Dict[str, Any]) -> str:
        question = self._resolve_question(session_result)
        observe = self._phase_result(session_result, "observe")
        literature = get_phase_value(observe, "literature_pipeline", {}) if isinstance(get_phase_value(observe, "literature_pipeline", {}), dict) else {}
        observations = self._coerce_string_list(get_phase_value(observe, "observations", []))
        findings = self._coerce_string_list(get_phase_value(observe, "findings", []))
        literature_summaries = self._extract_literature_summaries(literature)
        evidence_points = self._extract_evidence_points(literature)

        paragraphs = [
            (
                f"本研究围绕“{question}”展开，目标是在中医药理论、古籍文献与现代证据之间建立可复核的分析链条。"
                "中医科研的关键难点并不只是发现方药、证候或靶点，而是如何把零散文献、临床观察与知识建模结果组织成可追踪、可解释、可转化的学术叙事。"
                "因此，本报告采用 IMRD 结构，将研究问题、证据基础、方法路径、主要发现与讨论建议整合为标准化交付物。"
            ),
            (
                "在引言部分，我们重点回答三个问题：第一，为什么该研究问题具有理论与临床意义；第二，现有文献已经提供了哪些线索；第三，目前仍存在哪些研究空白需要通过后续研究设计进一步验证。"
                f"基于当前 session 的观察阶段输出，系统共提取 {len(observations)} 条观察线索、{len(findings)} 条初步发现，并对文献综述结果进行了结构化整理。"
            ),
        ]

        if literature_summaries:
            paragraphs.append(
                "文献综述提示，已有研究主要集中在以下方向："
                + "；".join(literature_summaries[:4])
                + "。这些线索说明，该主题具有一定的研究积累，但在证据整合、机制解释与方法一致性方面仍存在提升空间。"
            )
        else:
            paragraphs.append(
                "当前输入结果尚未给出完整的文献摘要列表，因此引言对文献综述采用保守写法。按照中医科研常见场景，相关研究通常涉及古籍理论溯源、现代实验验证、临床疗效评价与证据综合四个层面。"
            )

        if evidence_points:
            paragraphs.append(
                "从证据矩阵来看，现有研究主要提示："
                + "；".join(evidence_points[:4])
                + "。这说明研究问题并非完全空白，而是处于“已有证据但尚未系统整合”的阶段，适合通过结构化研究流程进一步明确核心结论。"
            )

        if observations or findings:
            paragraphs.append(
                "结合观察阶段的直接输出，研究背景还可以概括为："
                + "；".join((observations + findings)[:5])
                + "。基于这些背景，本研究进一步提出核心科学问题，即如何在方药配伍规律、证候机制与证据质量之间建立更加一致的解释框架。"
            )

        paragraphs.append(
            "因此，本报告的引言结论是：该研究问题兼具理论价值与现实意义，具有形成规范化研究设计和学术论文初稿的条件；后续章节将说明采用何种方法开展分析、获得了哪些结果，以及这些结果在现有文献语境中应当如何解读。"
        )
        return "\n\n".join(paragraphs)

    def _build_methods(self, session_result: Dict[str, Any]) -> str:
        question = self._resolve_question(session_result)
        observe = self._phase_result(session_result, "observe")
        experiment = self._phase_result(session_result, "experiment")
        analyze = self._phase_result(session_result, "analyze")
        protocol = self._extract_protocol(experiment)
        data_sources = self._extract_data_sources(observe)
        analysis_methods = self._extract_analysis_methods(analyze)

        design_type = self._resolve_string(protocol.get("study_type")) or "综合性中医科研设计"
        sample_size = protocol.get("sample_size") if isinstance(protocol.get("sample_size"), dict) else {}
        estimated_n = sample_size.get("estimated_n", 0)
        pico = protocol.get("pico") if isinstance(protocol.get("pico"), dict) else {}
        eligibility = protocol.get("eligibility") if isinstance(protocol.get("eligibility"), dict) else {}

        paragraphs = [
            (
                f"方法学部分围绕“{question}”组织，整体采用 {design_type} 作为主线，并将文献检索、文本预处理、实体抽取、知识图谱建模与推理分析串联为闭环。"
                "这一设计的目的，是让中医科研过程从单点工具调用升级为阶段化研究流程，使研究者能够同时看到数据来源、处理逻辑与结论生成依据。"
            ),
            (
                "数据来源方面，本次研究优先整合古籍文本、结构化文献记录、观察阶段产生的摘要与证据矩阵，同时保留后续扩展实验、随访数据或外部数据库接入的接口。"
                + ("本次 session 已识别的数据来源包括：" + "、".join(data_sources[:6]) + "。" if data_sources else "当前结果未完整列出数据源名称，因此方法部分使用通用数据来源描述。")
            ),
        ]

        if pico:
            paragraphs.append(
                "为保证研究问题具备可执行性，实验设计尽量按 PICO 框架组织："
                f"研究对象为“{self._resolve_string(pico.get('population')) or '待明确人群'}”，"
                f"干预措施为“{self._resolve_string(pico.get('intervention')) or '待明确干预'}”，"
                f"对照方案为“{self._resolve_string(pico.get('comparison')) or '待明确对照'}”，"
                f"主要结局为“{self._resolve_string(pico.get('outcome')) or '待明确结局'}”。"
                "这一框架有助于后续开展系统综述、随机对照试验或机制验证时保持术语与指标的一致性。"
            )

        if estimated_n:
            paragraphs.append(
                f"样本量方面，当前协议给出的估算总样本量约为 {estimated_n} 例。"
                "虽然该估算仍需结合效应量、失访率与中心间差异进一步校正，但已经能够为正式研究启动提供数量级参考。"
            )

        if eligibility:
            inclusion = self._coerce_string_list(eligibility.get("inclusion"))
            exclusion = self._coerce_string_list(eligibility.get("exclusion"))
            paragraphs.append(
                f"纳排标准模板方面，当前协议设置了 {len(inclusion)} 条纳入标准与 {len(exclusion)} 条排除标准。"
                "这种模板化设计能够降低研究方案撰写时的遗漏风险，并提高伦理申报与统计分析计划书之间的一致性。"
            )

        if protocol:
            paragraphs.append(
                "需要说明的是，当前研究链中的 Experiment 阶段承担的是研究方案与验证协议设计职责，"
                "并不代表已经完成真实实验执行、临床入组或外部验证。"
                "因此，PICO、样本量、纳排标准与预期结局应被视为待执行 protocol 草案，后续仍需由研究者或外部执行系统落地。"
            )

        if analysis_methods:
            paragraphs.append(
                "分析方法方面，系统当前采用的核心路径包括："
                + "；".join(analysis_methods[:5])
                + "。这些方法共同服务于一个目标，即把中医知识发现从“描述性输出”推进到“可复核的结构化分析”。"
            )
        else:
            paragraphs.append(
                "分析方法默认包含文本清洗、命名实体抽取、语义关系建模、知识图谱统计与规则/证据驱动推理。"
                "若后续接入证据分级、研究设计辅助或论文写作模块，上述方法还可以进一步映射到 GRADE、PICO 与 IMRD 等标准框架。"
            )

        paragraphs.append(
            "综上，Methods 章节强调的不只是“做了哪些技术步骤”，而是这些技术步骤如何共同构成一个规范的科研闭环。"
            "对中医科研而言，这种结构化方法学表达有助于后续形成论文方法部分、研究方案草案以及项目申报材料。"
        )
        return "\n\n".join(paragraphs)

    def _build_results(self, session_result: Dict[str, Any]) -> str:
        observe = self._phase_result(session_result, "observe")
        analyze = self._phase_result(session_result, "analyze")
        entities = self._extract_entities(session_result)
        graph_summary = self._extract_graph_summary(session_result)
        reasoning_points = self._extract_reasoning_points(session_result)
        findings = self._coerce_string_list(get_phase_value(observe, "findings", []))
        analysis_points = self._extract_analysis_points(analyze)

        paragraphs = [
            (
                "结果部分聚焦于结构化发现、图谱化关系与推理性结论三个层次。首先，在实体发现层面，系统从文本和研究上下文中抽取了疾病、方剂、药物、证候、靶点或结局等关键对象，"
                f"当前可识别的代表性实体包括：{('、'.join(entities[:10])) if entities else '暂未提供实体清单'}。"
                "这些实体构成了后续知识建模与证据组织的基础。"
            ),
        ]

        if graph_summary["node_count"] or graph_summary["edge_count"]:
            paragraphs.append(
                f"在知识图谱层面，当前结果显示图谱大约包含 {graph_summary['node_count']} 个节点与 {graph_summary['edge_count']} 条边。"
                "节点用于表示方药、证候、机制与结局等对象，边用于表达配伍、关联、作用、证据支持等关系。"
                "图谱化表示的价值在于：它既能帮助研究者识别高频关系，也能为后续机制推断与假设生成提供结构约束。"
            )
        else:
            paragraphs.append(
                "在知识图谱层面，当前 session 尚未输出完整的节点和边统计，因此本节采取保守描述。"
                "但从系统设计上看，知识图谱承担了把离散研究事实组织成可计算网络的重要职责。"
            )

        if findings or analysis_points:
            paragraphs.append(
                "从观察与分析阶段的直接输出看，主要发现包括："
                + "；".join((findings + analysis_points)[:6])
                + "。这些发现提示，研究主题在临床应用逻辑、理论解释路径与证据一致性方面已经具有可形成论文结果段落的基础。"
            )

        if reasoning_points:
            paragraphs.append(
                "推理模块进一步给出的结论为："
                + "；".join(reasoning_points[:5])
                + "。与单纯罗列实体不同，这些结论强调不同对象之间的结构关系及其学术含义，因此更接近论文 Results 章节中“主结果”的表达方式。"
            )
        else:
            paragraphs.append(
                "在推理层面，若当前尚未输出明确结论，仍可依据实体频次、关系密度与观察线索形成初步研究发现。"
                "这意味着结果章节不仅可以呈现最终结论，也可以诚实展示“已识别的模式”与“仍待验证的关系”。"
            )

        paragraphs.append(
            "总体而言，本次 session 已经完成了从数据采集到结构化结果表达的核心转换：文本不再只是原始材料，而被转写为可用于学术写作的研究证据单元。"
            "这为后续讨论章节中的文献比较、局限性分析和未来研究方向提出提供了直接依据。"
        )
        return "\n\n".join(paragraphs)

    def _build_discussion(self, session_result: Dict[str, Any]) -> str:
        question = self._resolve_question(session_result)
        observe = self._phase_result(session_result, "observe")
        analyze = self._phase_result(session_result, "analyze")
        literature = get_phase_value(observe, "literature_pipeline", {}) if isinstance(get_phase_value(observe, "literature_pipeline", {}), dict) else {}
        literature_summaries = self._extract_literature_summaries(literature)
        limitations = self._extract_limitations(session_result)
        future_directions = self._extract_future_directions(session_result)
        comparison_points = self._extract_comparison_points(analyze)

        paragraphs = [
            (
                f"Discussion 章节围绕“{question}”对结果进行解释。"
                "从当前输出看，系统识别出的实体关系、图谱结构与推理结论，整体上与中医科研中常见的“理论线索先行、证据逐步补强”的知识演化模式相一致。"
                "这说明闭环式研究流程并非仅用于自动化运行，也能帮助研究者更清楚地界定结论边界。"
            ),
        ]

        if comparison_points:
            paragraphs.append(
                "与现有文献相比，当前结果可归纳出以下差异或补充："
                + "；".join(comparison_points[:4])
                + "。这类对比有助于区分“重复已有结论”与“在既有研究基础上进行了结构化推进”两种不同的学术贡献。"
            )
        elif literature_summaries:
            paragraphs.append(
                "与现有研究进行对照后可以发现，文献更多集中于疗效描述或单点机制解释，而当前报告尝试把文献综述、结构化实体、图谱关系与推理结果统一到一条证据链中。"
                "这一点意味着本研究更接近“可生成论文讨论段落的系统性分析”，而不是简单的摘要拼接。"
            )

        paragraphs.append(
            "不过，需要注意的是，自动化科研流程生成的结果并不等同于最终定稿结论。"
            "讨论部分必须保留对证据质量、样本代表性、知识抽取误差、关系推断偏差和外部验证不足的审慎判断。"
            "尤其是在中医复杂系统研究中，方药配伍规律往往受到病种、证型、剂量、疗程和数据来源差异的共同影响。"
        )

        if limitations:
            paragraphs.append(
                "当前阶段可识别的主要局限性包括："
                + "；".join(limitations[:5])
                + "。这些局限提醒我们，在正式对外发布前，仍需通过人工复核、补充实验或扩大证据覆盖范围来提高结论可信度。"
            )
        else:
            paragraphs.append(
                "若从系统方法角度总结局限性，至少包括三点：第一，输入数据的完整性会直接影响实体与关系抽取质量；第二，推理链条仍依赖已有规则和证据模板；第三，当前结论尚未经过前瞻性实验或高等级证据全面验证。"
            )

        if future_directions:
            paragraphs.append(
                "后续研究方向建议优先考虑："
                + "；".join(future_directions[:5])
                + "。这些方向能够帮助研究从“自动化分析”进一步走向“可发表、可复现、可转化”的完整学术成果。"
            )
        else:
            paragraphs.append(
                "未来工作可从三个方向推进：其一，补充高质量临床研究与系统综述证据；其二，结合实验设计辅助模块形成正式研究方案；其三，将 IMRD 报告与论文写作、图表生成和证据分级模块联动，形成更完整的交付链条。"
            )

        paragraphs.append(
            "综合来看，本研究报告已经具备论文 Discussion 章节的基本要素：能够解释结果、联系文献、揭示局限并提出未来方向。"
            "对中医科研数字化场景而言，这种输出方式能够显著缩短从“研究运行完成”到“学术报告成稿”的距离。"
        )
        return "\n\n".join(paragraphs)

    def _render_markdown(
        self,
        title: str,
        session_result: Dict[str, Any],
        sections: Dict[str, str],
    ) -> str:
        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
        question = self._resolve_question(session_result)
        parts = [
            f"# {title}",
            f"\n生成时间：{generated_at}\n",
            f"研究问题：{question}\n",
            "## Introduction\n\n" + sections["introduction"],
            "## Methods\n\n" + sections["methods"],
            "## Results\n\n" + sections["results"],
            "## Discussion\n\n" + sections["discussion"],
        ]
        return "\n\n".join(parts).strip() + "\n"

    def _render_markdown_file(self, session_result: Dict[str, Any], markdown_content: str) -> str:
        self.output_dir.mkdir(parents=True, exist_ok=True)
        output_path = self.output_dir / f"{self._resolve_file_stem(session_result)}_imrd_report.md"
        output_path.write_text(markdown_content, encoding="utf-8")
        return str(output_path)

    def _ensure_minimum_length(
        self,
        markdown_content: str,
        session_result: Dict[str, Any],
        sections: Dict[str, str],
    ) -> str:
        if len(markdown_content) >= self.min_markdown_length:
            return markdown_content
        appendix = (
            "\n\n## Appendix\n\n"
            "本附录用于在科研 session 结果较少时补充结构化上下文，确保报告仍具备可审阅性。"
            f"当前研究问题为“{self._resolve_question(session_result)}”，"
            f"四个 IMRD 章节已全部生成，章节长度分别为：{json.dumps({k: len(v) for k, v in sections.items()}, ensure_ascii=False)}。"
            "研究者在正式发表前，可继续补充更多文献证据、实验设计细节和外部验证结果，以提高报告的完整性与学术说服力。"
        )
        return markdown_content + appendix

    def _render_docx(
        self,
        title: str,
        session_result: Dict[str, Any],
        sections: Dict[str, str],
        markdown_content: str,
    ) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("未检测到 python-docx，无法生成 DOCX 报告") from exc

        self.output_dir.mkdir(parents=True, exist_ok=True)
        output_path = self.output_dir / f"{self._resolve_file_stem(session_result)}_imrd_report.docx"
        document = Document()
        document.add_heading(title, level=0)
        document.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        document.add_paragraph(f"研究问题：{self._resolve_question(session_result)}")

        for heading, key in (
            ("Introduction", "introduction"),
            ("Methods", "methods"),
            ("Results", "results"),
            ("Discussion", "discussion"),
        ):
            document.add_heading(heading, level=1)
            for paragraph in self._split_paragraphs(sections[key]):
                document.add_paragraph(paragraph)

        document.add_heading("Appendix", level=1)
        document.add_paragraph(f"Markdown 字符数：{len(markdown_content)}")
        document.save(output_path)
        return str(output_path)

    def _extract_literature_summaries(self, literature: Dict[str, Any]) -> List[str]:
        summaries = self._coerce_string_list(literature.get("summaries"))
        if summaries:
            return summaries
        records = literature.get("records")
        extracted: List[str] = []
        if isinstance(records, list):
            for record in records:
                if isinstance(record, dict):
                    title = self._resolve_string(record.get("title"))
                    abstract = self._resolve_string(record.get("abstract"))
                    if title and abstract:
                        extracted.append(f"{title}：{abstract[:60]}")
                    elif title:
                        extracted.append(title)
        return extracted

    def _extract_evidence_points(self, literature: Dict[str, Any]) -> List[str]:
        matrix = literature.get("evidence_matrix")
        points: List[str] = []
        if isinstance(matrix, list):
            for item in matrix:
                if isinstance(item, dict):
                    intervention = self._resolve_string(item.get("intervention"))
                    outcome = self._resolve_string(item.get("outcome"))
                    evidence = self._resolve_string(item.get("evidence_level")) or self._resolve_string(item.get("grade"))
                    segment = " - ".join(part for part in (intervention, outcome, evidence) if part)
                    if segment:
                        points.append(segment)
                else:
                    text = self._resolve_string(item)
                    if text:
                        points.append(text)
        return points

    def _extract_protocol(self, experiment: Dict[str, Any]) -> Dict[str, Any]:
        for key in ("study_protocol", "protocol", "experiment_protocol"):
            candidate = get_phase_value(experiment, key)
            if isinstance(candidate, dict):
                return candidate
        experiments = get_phase_value(experiment, "experiments", [])
        if isinstance(experiments, list) and experiments:
            first_experiment = experiments[0]
            if isinstance(first_experiment, dict):
                nested_protocol = first_experiment.get("study_protocol")
                if isinstance(nested_protocol, dict):
                    return nested_protocol
        return {}

    def _extract_data_sources(self, observe: Dict[str, Any]) -> List[str]:
        sources: List[str] = []
        corpus = get_phase_value(observe, "corpus_collection", {}) if isinstance(get_phase_value(observe, "corpus_collection", {}), dict) else {}
        literature = get_phase_value(observe, "literature_pipeline", {}) if isinstance(get_phase_value(observe, "literature_pipeline", {}), dict) else {}
        ingestion = get_phase_value(observe, "ingestion_pipeline", {}) if isinstance(get_phase_value(observe, "ingestion_pipeline", {}), dict) else {}

        if corpus:
            sources.append("古籍/本地语料")
        records = literature.get("records")
        if isinstance(records, list) and records:
            source_names = []
            for record in records:
                if isinstance(record, dict):
                    source_name = self._resolve_string(record.get("source"))
                    if source_name:
                        source_names.append(source_name)
            if source_names:
                sources.extend(source_names)
            else:
                sources.append(f"结构化文献记录 {len(records)} 条")
        if literature:
            sources.append("文献摘要与证据矩阵")
        if ingestion:
            sources.append("预处理与语义建模结果")
        return self._unique_strings(sources)

    def _extract_analysis_methods(self, analyze: Dict[str, Any]) -> List[str]:
        methods = self._coerce_string_list(get_phase_value(analyze, "analysis_methods", []))
        if methods:
            return methods

        metadata = analyze.get("metadata") if isinstance(analyze.get("metadata"), dict) else {}
        collected = [
            _ANALYZE_METHOD_LABELS.get(str(item), str(item))
            for item in self._coerce_string_list(metadata.get("analysis_modules"))
        ]
        collected.extend(
            _ANALYZE_METHOD_LABELS.get(str(item), str(item))
            for item in self._coerce_string_list(metadata.get("data_mining_methods"))
        )
        if collected:
            return self._unique_strings(collected)

        collected = []
        if analyze:
            collected.append("实体抽取与关系建模")
            collected.append("知识图谱统计")
            collected.append("规则/证据驱动推理")
        return self._unique_strings(collected)

    def _extract_analysis_points(self, analyze: Dict[str, Any]) -> List[str]:
        direct_points = self._coerce_string_list(get_phase_value(analyze, "analysis_results", []))
        if direct_points:
            return direct_points

        statistical_analysis = get_phase_value(analyze, "statistical_analysis", {})
        if not isinstance(statistical_analysis, dict):
            return []

        collected: List[str] = []
        interpretation = self._resolve_string(statistical_analysis.get("interpretation"))
        if interpretation:
            collected.append(interpretation)

        evidence_grade_summary = statistical_analysis.get("evidence_grade_summary")
        if isinstance(evidence_grade_summary, dict):
            collected.extend(self._coerce_string_list(evidence_grade_summary.get("summary")))
        return self._unique_strings(collected)

    def _extract_comparison_points(self, analyze: Dict[str, Any]) -> List[str]:
        return self._coerce_string_list(get_phase_value(analyze, "comparison_with_literature", []))

    def _extract_entities(self, session_result: Dict[str, Any]) -> List[str]:
        raw_values = self._collect_by_keys(session_result, {"entities", "entity_list", "named_entities"})
        entities: List[str] = []
        for value in raw_values:
            if isinstance(value, list):
                for item in value:
                    if isinstance(item, dict):
                        name = self._resolve_string(item.get("name")) or self._resolve_string(item.get("entity"))
                        if name:
                            entities.append(name)
                    else:
                        name = self._resolve_string(item)
                        if name:
                            entities.append(name)
            elif isinstance(value, dict):
                name = self._resolve_string(value.get("name")) or self._resolve_string(value.get("entity"))
                if name:
                    entities.append(name)
            else:
                name = self._resolve_string(value)
                if name:
                    entities.append(name)
        return self._unique_strings(entities)

    def _extract_graph_summary(self, session_result: Dict[str, Any]) -> Dict[str, int]:
        graph_values = self._collect_by_keys(session_result, {"semantic_graph", "knowledge_graph", "graph"})
        for graph in graph_values:
            if isinstance(graph, dict):
                nodes = graph.get("nodes") or graph.get("entities") or []
                edges = graph.get("edges") or graph.get("relationships") or []
                node_count = len(nodes) if isinstance(nodes, list) else int(graph.get("node_count", 0) or 0)
                edge_count = len(edges) if isinstance(edges, list) else int(graph.get("edge_count", 0) or 0)
                if node_count or edge_count:
                    return {"node_count": node_count, "edge_count": edge_count}
        return {"node_count": 0, "edge_count": 0}

    def _extract_reasoning_points(self, session_result: Dict[str, Any]) -> List[str]:
        raw_values = self._collect_by_keys(
            session_result,
            {"reasoning_results", "reasoning_conclusions", "insights", "conclusions", "analysis_findings"},
        )
        points: List[str] = []
        for value in raw_values:
            if isinstance(value, list):
                for item in value:
                    if isinstance(item, dict):
                        text = self._join_non_empty(
                            self._resolve_string(item.get("title")),
                            self._resolve_string(item.get("description")),
                            separator="：",
                        )
                        if text:
                            points.append(text)
                    else:
                        text = self._resolve_string(item)
                        if text:
                            points.append(text)
            elif isinstance(value, dict):
                text = self._join_non_empty(
                    self._resolve_string(value.get("title")),
                    self._resolve_string(value.get("description")),
                    separator="：",
                )
                if text:
                    points.append(text)
            else:
                text = self._resolve_string(value)
                if text:
                    points.append(text)
        return self._unique_strings(points)

    def _extract_limitations(self, session_result: Dict[str, Any]) -> List[str]:
        reflect = self._phase_result(session_result, "reflect")
        analyze = self._phase_result(session_result, "analyze")
        limitations = self._coerce_string_list(get_phase_value(reflect, "limitations", []))
        limitations.extend(self._extract_analyze_limitations(analyze))
        return self._unique_strings(limitations)

    def _extract_future_directions(self, session_result: Dict[str, Any]) -> List[str]:
        reflect = self._phase_result(session_result, "reflect")
        publish = self._phase_result(session_result, "publish")
        directions = self._coerce_string_list(get_phase_value(reflect, "improvement_plan", []))
        directions.extend(self._coerce_string_list(get_phase_value(reflect, "future_directions", [])))
        directions.extend(self._coerce_string_list(get_phase_value(reflect, "recommendations", [])))

        publish_analysis_results = get_phase_value(publish, "analysis_results", {})
        if isinstance(publish_analysis_results, dict):
            directions.extend(self._coerce_string_list(publish_analysis_results.get("recommendations")))
        directions.extend(self._coerce_string_list(get_phase_value(publish, "recommendations", [])))
        return self._unique_strings(directions)

    def _extract_analyze_limitations(self, analyze: Dict[str, Any]) -> List[str]:
        direct_limitations = self._coerce_string_list(get_phase_value(analyze, "limitations", []))
        if direct_limitations:
            return direct_limitations

        statistical_analysis = get_phase_value(analyze, "statistical_analysis", {})
        if isinstance(statistical_analysis, dict):
            return self._coerce_string_list(statistical_analysis.get("limitations"))
        return []

    def _collect_by_keys(self, payload: Any, target_keys: set[str]) -> List[Any]:
        collected: List[Any] = []

        def visit(node: Any) -> None:
            if isinstance(node, dict):
                for key, value in node.items():
                    if key in target_keys:
                        collected.append(value)
                    visit(value)
            elif isinstance(node, list):
                for item in node:
                    visit(item)

        visit(payload)
        return collected

    def _split_paragraphs(self, content: str) -> List[str]:
        return [paragraph.strip() for paragraph in content.split("\n\n") if paragraph.strip()]

    def _coerce_string_list(self, value: Any) -> List[str]:
        if value is None:
            return []
        if isinstance(value, list):
            flattened: List[str] = []
            for item in value:
                if isinstance(item, dict):
                    text = self._resolve_string(item.get("description")) or self._resolve_string(item.get("title"))
                    if text:
                        flattened.append(text)
                else:
                    text = self._resolve_string(item)
                    if text:
                        flattened.append(text)
            return flattened
        text = self._resolve_string(value)
        return [text] if text else []

    def _resolve_string(self, value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, str):
            return value.strip()
        if isinstance(value, (int, float)):
            return str(value)
        return ""

    def _unique_strings(self, values: Iterable[str]) -> List[str]:
        seen = set()
        result: List[str] = []
        for value in values:
            normalized = value.strip()
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            result.append(normalized)
        return result

    def _join_non_empty(self, *parts: str, separator: str = " ") -> str:
        normalized = [part.strip() for part in parts if part and part.strip()]
        return separator.join(normalized)

    def _slugify(self, text: str) -> str:
        slug = re.sub(r"[^0-9A-Za-z_\-]+", "_", text).strip("_")
        return slug or f"report_{int(time.time())}"

    def _resolve_file_stem(self, session_result: Dict[str, Any]) -> str:
        session_id = self._resolve_string(session_result.get("session_id"))
        title = self._resolve_string(session_result.get("title")) or self._resolve_title(session_result)
        return self._slugify(session_id or title)


__all__ = ["Report", "ReportFormat", "ReportGenerator"]