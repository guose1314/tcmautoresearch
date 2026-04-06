# -*- coding: utf-8 -*-
"""论文写作辅助 — IMRD 大纲生成、模板填充、参考文献格式化与 DOCX 导出。"""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# IMRD section definitions
# ---------------------------------------------------------------------------

_IMRD_SECTIONS = ["introduction", "methods", "results", "discussion"]
_SECTION_TITLES_ZH = {
    "introduction": "引言（Introduction）",
    "methods": "方法（Methods）",
    "results": "结果（Results）",
    "discussion": "讨论（Discussion）",
}

_OUTLINE_SYSTEM = (
    "你是一位中医药学术论文写作专家。请根据研究主题和已有数据，"
    "生成一份 IMRD 结构的论文大纲。返回严格 JSON 对象，包含：\n"
    "title — 论文标题\n"
    "abstract_outline — 摘要要点(数组)\n"
    "keywords — 关键词(数组)\n"
    "introduction — {background, gap, objective, significance}(各为字符串)\n"
    "methods — {study_design, subjects, interventions, outcome_measures, statistical_analysis}(各为字符串)\n"
    "results — {primary_outcomes, secondary_outcomes, subgroup_analysis}(各为字符串)\n"
    "discussion — {main_findings, comparison_with_literature, limitations, future_directions}(各为字符串)\n"
    "不要添加 JSON 以外的文字。"
)

_FILL_SYSTEM = (
    "你是一位中医药学术论文撰写助手。请根据提供的大纲和研究数据，"
    "为论文各节生成完整的学术段落。返回 JSON 对象，每个键对应一节的内容字符串。"
)


class WritingHelper:
    """论文写作辅助引擎。

    Parameters
    ----------
    llm_engine : object | None
        LLM 推理引擎，需具备 ``generate(prompt, system_prompt)`` 方法。
    """

    def __init__(self, llm_engine: Optional[Any] = None) -> None:
        self._llm = llm_engine

    # ------------------------------------------------------------------
    # 1. Generate outline
    # ------------------------------------------------------------------

    def generate_outline(
        self,
        topic: str,
        research_data: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """生成 IMRD 论文大纲。

        Returns
        -------
        dict
            含 title, abstract_outline, keywords, introduction, methods,
            results, discussion 各节大纲要点。
        """
        parts = [f"【研究主题】\n{topic}"]
        if research_data:
            parts.append(f"【研究数据摘要】\n{json.dumps(research_data, ensure_ascii=False, default=str)[:3000]}")
        parts.append("请生成 IMRD 论文大纲（JSON 对象）。")
        prompt = "\n\n".join(parts)

        raw = self._call_llm(prompt, _OUTLINE_SYSTEM)
        outline = self._parse_json_dict(raw)

        if not outline or "introduction" not in outline:
            outline = self._fallback_outline(topic, research_data)

        return outline

    # ------------------------------------------------------------------
    # 2. Fill IMRD template
    # ------------------------------------------------------------------

    def fill_imrd_template(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """填充 IMRD 各节内容。

        Parameters
        ----------
        data : dict
            含 outline（大纲）和可选 research_data（研究数据）。

        Returns
        -------
        dict
            各节填充后的完整文本: ``{title, abstract, introduction, methods, results, discussion}``。
        """
        outline = data.get("outline", data)
        research_data = data.get("research_data", {})

        prompt_parts = [
            f"【论文大纲】\n{json.dumps(outline, ensure_ascii=False, default=str)[:4000]}",
        ]
        if research_data:
            prompt_parts.append(f"【研究数据】\n{json.dumps(research_data, ensure_ascii=False, default=str)[:3000]}")
        prompt_parts.append("请为 Introduction、Methods、Results、Discussion 各节撰写完整学术段落（JSON 对象）。")
        prompt = "\n\n".join(prompt_parts)

        raw = self._call_llm(prompt, _FILL_SYSTEM)
        filled = self._parse_json_dict(raw)

        # 确保必需字段
        result: Dict[str, Any] = {
            "title": outline.get("title", "未命名论文"),
            "abstract": filled.get("abstract", ""),
            "keywords": outline.get("keywords", []),
        }
        for section in _IMRD_SECTIONS:
            result[section] = filled.get(section, self._section_placeholder(section, outline))

        return result

    # ------------------------------------------------------------------
    # 3. Format references
    # ------------------------------------------------------------------

    @staticmethod
    def format_references(
        refs: List[Dict[str, Any]],
        style: str = "gb_t_7714",
    ) -> List[str]:
        """将参考文献列表格式化为指定引用风格。

        Parameters
        ----------
        refs : list[dict]
            每项可含 authors, title, journal, year, volume, issue, pages, doi。
        style : str
            引用风格，当前支持 ``"gb_t_7714"``（GB/T 7714-2015）。

        Returns
        -------
        list[str]
            格式化后的参考文献字符串列表。
        """
        formatted: List[str] = []
        for idx, ref in enumerate(refs, 1):
            if style == "gb_t_7714":
                formatted.append(_format_gb_t_7714(idx, ref))
            else:
                formatted.append(_format_gb_t_7714(idx, ref))
        return formatted

    # ------------------------------------------------------------------
    # 4. Export DOCX
    # ------------------------------------------------------------------

    def export_docx(self, content: Dict[str, Any], output_path: str) -> str:
        """导出为 DOCX 文件。

        Parameters
        ----------
        content : dict
            fill_imrd_template 的返回结果。
        output_path : str
            输出文件路径。

        Returns
        -------
        str
            实际写入的文件绝对路径。
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        doc = Document()

        # 标题
        title = content.get("title", "未命名论文")
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 摘要
        abstract = content.get("abstract", "")
        if abstract:
            doc.add_heading("摘要", level=1)
            doc.add_paragraph(abstract)

        # 关键词
        keywords = content.get("keywords", [])
        if keywords:
            kw_para = doc.add_paragraph()
            run = kw_para.add_run("关键词：")
            run.bold = True
            kw_para.add_run("；".join(str(k) for k in keywords))

        doc.add_page_break()

        # IMRD 各节
        for idx, section in enumerate(_IMRD_SECTIONS, 1):
            section_title = f"{idx} {_SECTION_TITLES_ZH.get(section, section)}"
            doc.add_heading(section_title, level=1)
            text = content.get(section, "")
            if isinstance(text, dict):
                for sub_key, sub_val in text.items():
                    doc.add_heading(str(sub_key), level=2)
                    doc.add_paragraph(str(sub_val))
            else:
                doc.add_paragraph(str(text))

        # 参考文献
        references = content.get("references", [])
        if references:
            doc.add_page_break()
            doc.add_heading("参考文献", level=1)
            for ref_text in references:
                doc.add_paragraph(str(ref_text), style="List Number")

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font.size = Pt(12)

        # 写文件
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        logger.info("DOCX 已导出: %s", out.resolve())
        return str(out.resolve())

    # ------------------------------------------------------------------
    # LLM helpers
    # ------------------------------------------------------------------

    def _call_llm(self, prompt: str, system_prompt: str) -> str:
        engine = self._get_llm()
        if engine is None:
            return ""
        try:
            return engine.generate(prompt, system_prompt=system_prompt)
        except Exception:
            logger.exception("LLM 生成失败")
            return ""

    def _get_llm(self):
        if self._llm is not None:
            return self._llm
        try:
            from src.llm.llm_engine import LLMEngine
            engine = LLMEngine()
            engine.load()
            self._llm = engine
            return engine
        except Exception as exc:
            self._llm = None
            logger.warning("无法加载 LLM 引擎: %s", exc)
            return None

    # ------------------------------------------------------------------
    # JSON parsing
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_json_dict(text: str) -> Dict[str, Any]:
        if not text:
            return {}
        try:
            parsed = json.loads(text)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            pass
        m = re.search(r'```(?:json)?\s*(\{.*?})\s*```', text, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(1))
            except json.JSONDecodeError:
                pass
        m = re.search(r'\{.*}', text, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(0))
            except json.JSONDecodeError:
                pass
        return {}

    # ------------------------------------------------------------------
    # Fallbacks & placeholders
    # ------------------------------------------------------------------

    @staticmethod
    def _fallback_outline(topic: str, research_data: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        return {
            "title": f"{topic}的研究",
            "abstract_outline": ["研究背景与目的", "方法概述", "主要结果", "结论与意义"],
            "keywords": [topic, "中医药", "研究"],
            "introduction": {
                "background": f"{topic}的研究背景",
                "gap": "当前研究不足",
                "objective": f"本研究旨在探讨{topic}",
                "significance": "研究意义",
            },
            "methods": {
                "study_design": "待确定",
                "subjects": "待确定",
                "interventions": "待确定",
                "outcome_measures": "待确定",
                "statistical_analysis": "待确定",
            },
            "results": {
                "primary_outcomes": "待填充",
                "secondary_outcomes": "待填充",
                "subgroup_analysis": "待填充",
            },
            "discussion": {
                "main_findings": "待撰写",
                "comparison_with_literature": "待撰写",
                "limitations": "待撰写",
                "future_directions": "待撰写",
            },
        }

    @staticmethod
    def _section_placeholder(section: str, outline: Dict[str, Any]) -> str:
        data = outline.get(section, {})
        if isinstance(data, dict):
            return "\n\n".join(f"**{k}**: {v}" for k, v in data.items() if v)
        return str(data) if data else f"[{_SECTION_TITLES_ZH.get(section, section)} 待撰写]"


# ---------------------------------------------------------------------------
# GB/T 7714-2015 reference formatter
# ---------------------------------------------------------------------------

def _format_gb_t_7714(idx: int, ref: Dict[str, Any]) -> str:
    """格式化单条参考文献为 GB/T 7714-2015 格式。"""
    parts: List[str] = [f"[{idx}]"]

    # 作者
    authors = ref.get("authors", [])
    if isinstance(authors, str):
        authors = [a.strip() for a in authors.split(",")]
    if authors:
        author_str = ", ".join(authors[:3])
        if len(authors) > 3:
            author_str += ", 等"
        parts.append(f" {author_str}.")

    # 标题
    title = ref.get("title", "")
    if title:
        parts.append(f" {title}[J].")

    # 期刊
    journal = ref.get("journal", "")
    if journal:
        parts.append(f" {journal},")

    # 年份
    year = ref.get("year", "")
    if year:
        parts.append(f" {year}")

    # 卷期页
    volume = ref.get("volume", "")
    issue = ref.get("issue", "")
    pages = ref.get("pages", "")
    if volume:
        vol_str = str(volume)
        if issue:
            vol_str += f"({issue})"
        parts.append(f", {vol_str}")
    if pages:
        parts.append(f": {pages}.")
    elif parts[-1][-1] != ".":
        parts.append(".")

    # DOI
    doi = ref.get("doi", "")
    if doi:
        parts.append(f" DOI: {doi}.")

    return "".join(parts)
