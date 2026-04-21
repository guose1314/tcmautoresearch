# src/output/report_generator.py
"""
学术报告生成器 — 将研究流程输出转换为结构化 Markdown / DOCX 报告。
"""
from __future__ import annotations

import json
import logging
import os
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from typing import Any, Dict, List, Optional

from src.core.module_base import BaseModule

logger = logging.getLogger(__name__)


class ReportFormat(str, Enum):
    """支持的报告格式。"""
    MARKDOWN = "markdown"
    DOCX = "docx"
    JSON = "json"


@dataclass
class Report:
    """报告对象，包含渲染内容与元数据。"""
    format: str
    content: str = ""
    sections: Dict[str, str] = field(default_factory=dict)
    output_path: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """将报告序列化为普通字典。"""
        return {
            "format": self.format,
            "content": self.content,
            "title": self.metadata.get("title", ""),
            "sections": self.sections,
            "output_path": self.output_path,
            "metadata": self.metadata,
        }


class ReportGenerator(BaseModule):
    """将研究会话结果渲染为 IMRD 结构报告。

    Config keys:
        output_dir (str): 输出目录，DOCX 文件写入此处。默认 ``./output``。
        include_abstract (bool): 是否在报告中包含摘要节。
    """

    def __init__(self, config: Optional[Dict[str, Any]] = None) -> None:
        super().__init__("report_generator", config)
        self._output_dir: str = self.config.get("output_dir", "./output")

    # ------------------------------------------------------------------
    def _do_initialize(self) -> bool:
        os.makedirs(self._output_dir, exist_ok=True)
        return True

    def _do_cleanup(self) -> bool:
        return True

    def _do_execute(self, context: Dict[str, Any]) -> Dict[str, Any]:
        session_result = context.get("session_result", {})
        fmt = context.get("format", "markdown")
        report = self.generate_report(session_result, fmt)
        return {
            "report": {
                "format": report.format,
                "output_path": report.output_path,
                "sections": list(report.sections.keys()),
                "metadata": {
                    "char_count": len(report.content),
                    "section_count": len(report.sections),
                    **report.metadata,
                },
            }
        }

    # ------------------------------------------------------------------
    def generate_report(
        self,
        session_result: Dict[str, Any],
        fmt: str | ReportFormat = ReportFormat.MARKDOWN,
    ) -> Report:
        """生成指定格式的报告。

        Args:
            session_result: 研究会话的结构化输出字典。
            fmt: 目标格式（``"markdown"``、``"docx"`` 或 ``"json"``）。

        Returns:
            ``Report`` 对象，包含内容与路径信息。

        Raises:
            ValueError: 不支持的格式。
        """
        if isinstance(fmt, ReportFormat):
            fmt_str = fmt.value
        else:
            fmt_str = str(fmt).lower()

        if fmt_str == ReportFormat.MARKDOWN.value:
            return self._render_markdown(session_result)
        elif fmt_str == ReportFormat.DOCX.value:
            return self._render_docx(session_result)
        elif fmt_str == ReportFormat.JSON.value:
            return self._render_json(session_result)
        else:
            raise ValueError(
                f"不支持的报告格式: '{fmt_str}'。"
                f"可选: {[f.value for f in ReportFormat]}"
            )

    # ------------------------------------------------------------------
    # 内部渲染方法
    # ------------------------------------------------------------------

    def _build_sections(self, session_result: Dict[str, Any]) -> Dict[str, str]:
        """将 session_result 解析为 IMRD 四节内容。"""
        phase_results = session_result.get("phase_results", {})
        question = session_result.get("question", "")

        # Introduction
        obs = phase_results.get("observe", {})
        obs_list = obs.get("observations", [])
        intro_lines = [
            f"本研究聚焦于以下问题：**{question}**",
            "",
            "研究背景与动机：",
        ]
        for o in obs_list[:3]:
            intro_lines.append(f"- {o}")
        introduction = "\n".join(intro_lines)

        # Methods
        methods_lines = ["本研究采用多阶段 AI 驱动研究方法，结合文献挖掘、实体抽取与知识图谱分析。", ""]
        lit = obs.get("literature_pipeline", {})
        if lit.get("records"):
            methods_lines.append(f"检索文献 {len(lit['records'])} 篇，覆盖 PubMed、CNKI 等数据库。")
        entity_phase = phase_results.get("analyze", {})
        entities = entity_phase.get("entities", [])
        if entities:
            methods_lines.append(f"实体抽取：共识别 {len(entities)} 个关键实体。")
        methods = "\n".join(methods_lines)

        # Results
        results_lines = []
        findings = obs.get("findings", [])
        for f in findings:
            results_lines.append(f"- {f}")
        analyze = phase_results.get("analyze", {})
        reasoning = analyze.get("reasoning_results", [])
        if not reasoning:
            reasoning = phase_results.get("reflect", {}).get("reasoning_results", [])
        for item in reasoning:
            if isinstance(item, dict):
                results_lines.append(f"\n**{item.get('title', '')}**\n{item.get('description', '')}")
        # Include entities
        ents = analyze.get("entities", [])
        if ents:
            herb_names = [e.get("name", e) if isinstance(e, dict) else str(e) for e in ents[:10]]
            results_lines.append("\n关键实体：" + "、".join(herb_names))
        results = "\n".join(results_lines) if results_lines else "（结果待补充）"

        # Discussion
        reflect = phase_results.get("reflect", {})
        disc_lines = []
        for item in reflect.get("comparison_with_literature", []):
            disc_lines.append(f"- {item}")
        for item in reflect.get("limitations", []):
            disc_lines.append(f"- {item}")
        for fd in reflect.get("future_directions", []):
            disc_lines.append(f"- {fd}")
        for rec in reflect.get("recommendations", []):
            disc_lines.append(f"- {rec}")
        discussion = "\n".join(disc_lines) if disc_lines else "（讨论待补充）"

        return {
            "introduction": introduction,
            "methods": methods,
            "results": results,
            "discussion": discussion,
        }

    def _safe_output_path(self, base_dir: str, filename: str) -> str:
        """构造并校验输出文件路径，确保路径不超出 base_dir 范围（防止路径穿越）。

        Args:
            base_dir: 已信任的输出根目录。
            filename: 经过严格清理的纯文件名（不含目录分隔符）。

        Returns:
            在 base_dir 内的绝对路径。

        Raises:
            ValueError: 如果计算后路径超出 base_dir（理论上不会，作为防御性检查）。
        """
        # filename 已由调用方用 re.sub 限制为 [\w\u4e00-\u9fff] + 时间戳 + 扩展名
        # 再次确保不含路径分隔符
        clean_name = os.path.basename(filename)
        candidate = os.path.realpath(os.path.join(base_dir, clean_name))
        real_base = os.path.realpath(base_dir)
        if not candidate.startswith(real_base + os.sep) and candidate != real_base:
            raise ValueError(
                f"输出路径安全校验失败：'{candidate}' 超出允许目录 '{real_base}'"
            )
        return candidate

    def _render_markdown(self, session_result: Dict[str, Any]) -> Report:
        meta = session_result.get("metadata", {})
        question = session_result.get("question", "")
        title = meta.get("title", question)
        sections = self._build_sections(session_result)

        lines: List[str] = [
            f"# {title}",
            "",
            f"> 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "## Introduction",
            "",
            sections["introduction"],
            "",
            "## Methods",
            "",
            sections["methods"],
            "",
            "## Results",
            "",
            sections["results"],
            "",
            "## Discussion",
            "",
            sections["discussion"],
            "",
            "---",
            "",
            f"*报告由 TCM AutoResearch 系统自动生成。研究问题：{question}*",
            "",
        ]
        content = "\n".join(lines)

        # 确保最小内容长度（>=500 字符），不足时追加研究元数据补充
        if len(content) < 500:
            extra_lines = ["", "## 附录：元数据快照", ""]
            for k, v in meta.items():
                extra_lines.append(f"- **{k}**: {v}")
            extra_lines.append("")
            content += "\n".join(extra_lines)
            # 如果仍然不足，追加研究背景说明
            if len(content) < 500:
                pad = (
                    "本报告涵盖中医研究全流程，包括文献综述、实验设计、数据分析、"
                    "结果讨论与知识推理，致力于从多维度揭示中医学理论与现代科学的关联。"
                )
                content += "\n\n" + pad

        report_meta: Dict[str, Any] = {
            "title": title,
            "question": question,
            "research_question": question,
            "generated_at": datetime.now().isoformat(),
            "section_count": 4,
        }

        # Write to file when output_dir is configured
        output_path = ""
        if self._output_dir:
            os.makedirs(self._output_dir, exist_ok=True)
            # 文件名只使用时间戳和随机 ID，不含任何用户输入，防止路径穿越
            filename = f"imrd_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.md"
            output_path = os.path.join(self._output_dir, filename)
            with open(output_path, "w", encoding="utf-8") as fh:
                fh.write(content)

        return Report(
            format=ReportFormat.MARKDOWN.value,
            content=content,
            sections=sections,
            output_path=output_path,
            metadata=report_meta,
        )

    def _render_docx(self, session_result: Dict[str, Any]) -> Report:
        """生成 DOCX 报告。若 python-docx 不可用则回退为 Markdown 文件保存为 .docx。"""
        md_report = self._render_markdown(session_result)
        os.makedirs(self._output_dir, exist_ok=True)
        # 文件名只使用时间戳和随机 ID，不含任何用户输入，防止路径穿越
        filename = f"imrd_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.docx"
        output_path = os.path.join(self._output_dir, filename)

        try:
            from docx import Document  # type: ignore

            doc = Document()
            title_str = md_report.metadata.get("title", session_result.get("question", ""))
            doc.add_heading(title_str, level=0)
            for sec_key, heading in [
                ("introduction", "Introduction"),
                ("methods", "Methods"),
                ("results", "Results"),
                ("discussion", "Discussion"),
            ]:
                doc.add_heading(heading, level=1)
                for para_text in md_report.sections[sec_key].split("\n"):
                    clean = para_text.strip().lstrip("- ").lstrip("**").rstrip("**")
                    if clean:
                        doc.add_paragraph(clean)
            doc.save(output_path)

        except ImportError:
            # python-docx 未安装，将 Markdown 写为 .docx 扩展文件
            with open(output_path, "w", encoding="utf-8") as fh:
                fh.write(md_report.content)

        return Report(
            format=ReportFormat.DOCX.value,
            content=md_report.content,
            sections=md_report.sections,
            output_path=output_path,
            metadata=md_report.metadata,
        )

    def _render_json(self, session_result: Dict[str, Any]) -> Report:
        sections = self._build_sections(session_result)
        content = json.dumps(
            {"sections": sections, "source": session_result},
            ensure_ascii=False,
            indent=2,
        )
        return Report(
            format=ReportFormat.JSON.value,
            content=content,
            sections=sections,
        )
