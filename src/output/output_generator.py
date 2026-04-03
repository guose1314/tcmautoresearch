# src/output/output_generator.py
"""
输出生成器模块 — 将管线结果序列化为多种格式。

支持功能：
- 管线结果结构化输出（metadata / analysis_results / quality_metrics / recommendations）
- 源文件路径安全净化（防止路径遍历）
- 实体数量限制、字符串长度限制
- JSON 安全化（将不可序列化对象转为字符串）
- 文件写出：JSON / Markdown / DOCX
"""

from __future__ import annotations

import json
import logging
import os
from datetime import datetime
from typing import Any, Dict, List, Optional

from src.core.module_base import BaseModule

logger = logging.getLogger(__name__)

_MAX_JSON_DEPTH = 8


class OutputGenerator(BaseModule):
    """将管线处理结果输出为结构化数据或文件。"""

    SUPPORTED_FORMATS = ("json", "markdown", "docx")

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__("output_generator", config)
        self.default_format: str = self.config.get("default_format", "json")
        self.default_output_dir: str = self.config.get("output_dir", "output")
        self.max_entities: int = int(self.config.get("max_entities", 500))
        self.max_string_length: int = int(self.config.get("max_string_length", 10000))
        self.max_recommendations: int = int(self.config.get("max_recommendations", 10))

    # ------------------------------------------------------------------
    # BaseModule 生命周期
    # ------------------------------------------------------------------

    def _do_initialize(self) -> bool:
        try:
            os.makedirs(self.default_output_dir, exist_ok=True)
            self.logger.info("OutputGenerator 初始化完成, 输出目录: %s", self.default_output_dir)
            return True
        except Exception:
            return False

    def _do_execute(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """
        执行输出生成。

        支持两种模式：
        1. 文件写出模式：context 含 ``output_format`` 键（json/markdown/docx）
        2. 结构化输出模式：context 含 ``entities`` 等管线结果键
        """
        # Allow subclass hook (for extensibility / testing)
        self._generate_output_format(context)

        # 若上层指定了输出格式，走文件写出
        if "output_format" in context:
            return self._file_output(context)

        # 否则走结构化输出（兼容现有测试）
        return self._structured_output(context)

    def _generate_output_format(self, context: Dict[str, Any]) -> None:
        """子类可覆写的输出格式化钩子（默认无操作）。"""
        pass

    def _do_cleanup(self) -> bool:
        try:
            self.logger.info("OutputGenerator 清理完成")
            return True
        except Exception:
            return False

    # ------------------------------------------------------------------
    # 结构化输出（管线内处理）
    # ------------------------------------------------------------------

    def _structured_output(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """将管线 context 整理为安全的结构化输出。"""
        # 安全净化源文件路径
        source = os.path.basename(context.get("source_file", "unknown"))

        # 限制实体数量
        entities = context.get("entities", [])
        if len(entities) > self.max_entities:
            entities = entities[: self.max_entities]

        # 质量统计（非负化）
        stats = context.get("statistics", {})
        quality_metrics = {
            "formulas_found": max(0, int(stats.get("formulas_count", 0))),
            "herbs_identified": max(0, int(stats.get("herbs_count", 0))),
            "syndromes_recognized": max(0, int(stats.get("syndromes_count", 0))),
        }

        # 分析结果（JSON 安全化）
        reasoning = context.get("reasoning_results", {})
        analysis_results = {
            "entities": entities,
            "reasoning_results": self._make_json_safe(reasoning),
            "objective": self._truncate(context.get("objective", ""), self.max_string_length),
        }

        # 推荐
        recommendations = self._build_recommendations(context)

        output_data = {
            "metadata": {
                "source": source,
                "generated_at": datetime.now().isoformat(),
            },
            "analysis_results": analysis_results,
            "quality_metrics": quality_metrics,
            "recommendations": recommendations,
        }

        return {"output_data": output_data, "success": True}

    # ------------------------------------------------------------------
    # JSON 安全化
    # ------------------------------------------------------------------

    def _make_json_safe(self, data: Any, depth: int = 0) -> Any:
        """递归将不可 JSON 序列化的对象转为截断字符串。"""
        if depth > _MAX_JSON_DEPTH:
            return self._truncate(str(data), self.max_string_length)

        if isinstance(data, dict):
            return {k: self._make_json_safe(v, depth + 1) for k, v in data.items()}
        if isinstance(data, (list, tuple)):
            return [self._make_json_safe(item, depth + 1) for item in data]
        if isinstance(data, (str, int, float, bool, type(None))):
            if isinstance(data, str):
                return self._truncate(data, self.max_string_length)
            return data
        # 不可序列化对象 → 字符串
        return self._truncate(str(data), self.max_string_length)

    def _truncate(self, text: str, max_len: int) -> str:
        if len(text) <= max_len:
            return text
        return text[:max_len]

    # ------------------------------------------------------------------
    # 推荐生成
    # ------------------------------------------------------------------

    def _build_recommendations(self, context: Dict[str, Any]) -> List[str]:
        """根据 context 生成改进推荐列表。"""
        recs: List[str] = []
        entities = context.get("entities", [])
        confidence = context.get("confidence_score", 0.0)

        if len(entities) < 5:
            recs.append("实体数量较少，建议扩大文献检索范围")
        if len(entities) >= 50:
            recs.append("实体数量丰富，建议聚焦核心主题实体")
        if confidence < 0.7:
            recs.append("置信度偏低，建议增加验证来源")
        if confidence >= 0.9:
            recs.append("置信度较高，可进入深度分析阶段")

        if not recs:
            recs.append("结果正常，无额外建议")

        return recs[: self.max_recommendations]

    # ------------------------------------------------------------------
    # 文件写出模式
    # ------------------------------------------------------------------

    def _file_output(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """写出文件并返回路径信息。"""
        pipeline_results = context.get("pipeline_results", context)
        output_format = context.get("output_format", self.default_format)
        output_dir = context.get("output_dir", self.default_output_dir)
        prefix = context.get("filename_prefix", "research_output")

        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if output_format == "markdown":
            path = self._write_markdown(pipeline_results, output_dir, prefix, timestamp)
        elif output_format == "docx":
            path = self._write_docx(pipeline_results, output_dir, prefix, timestamp)
        else:
            path = self._write_json(pipeline_results, output_dir, prefix, timestamp)

        self.logger.info("输出已写入: %s", path)
        return {"output_path": path, "output_format": output_format, "success": True}

    def _write_json(
        self, data: Dict[str, Any], output_dir: str, prefix: str, timestamp: str,
    ) -> str:
        path = os.path.join(output_dir, f"{prefix}_{timestamp}.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2, default=str)
        return path

    def _write_markdown(
        self, data: Dict[str, Any], output_dir: str, prefix: str, timestamp: str,
    ) -> str:
        path = os.path.join(output_dir, f"{prefix}_{timestamp}.md")
        lines: List[str] = [f"# 研究输出报告\n", f"生成时间: {timestamp}\n", "---\n"]
        self._dict_to_markdown(data, lines, level=2)
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return path

    def _dict_to_markdown(self, data: Any, lines: List[str], level: int = 2) -> None:
        if isinstance(data, dict):
            for key, value in data.items():
                lines.append(f"{'#' * level} {key}\n")
                self._dict_to_markdown(value, lines, level + 1)
        elif isinstance(data, list):
            for item in data:
                if isinstance(item, dict):
                    self._dict_to_markdown(item, lines, level)
                else:
                    lines.append(f"- {item}")
            lines.append("")
        else:
            lines.append(f"{data}\n")

    def _write_docx(
        self, data: Dict[str, Any], output_dir: str, prefix: str, timestamp: str,
    ) -> str:
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError:
            self.logger.warning("python-docx 未安装，降级为 JSON 格式输出")
            return self._write_json(data, output_dir, prefix, timestamp)
        path = os.path.join(output_dir, f"{prefix}_{timestamp}.docx")
        doc = Document()
        doc.add_heading("研究输出报告", level=0)
        doc.add_paragraph(f"生成时间: {timestamp}")
        self._dict_to_docx(doc, data, level=1)
        doc.save(path)
        return path

    def _dict_to_docx(self, doc: Any, data: Any, level: int = 1) -> None:
        if isinstance(data, dict):
            for key, value in data.items():
                doc.add_heading(str(key), level=min(level, 9))
                self._dict_to_docx(doc, value, level + 1)
        elif isinstance(data, list):
            for item in data:
                if isinstance(item, dict):
                    self._dict_to_docx(doc, item, level)
                else:
                    doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph(str(data))
