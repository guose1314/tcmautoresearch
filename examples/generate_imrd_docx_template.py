"""
使用 python-docx 生成 IMRD（Introduction, Methods, Results, Discussion）论文模板。

示例：
  c:/Users/hgk/tcmautoresearch/venv310/Scripts/python.exe examples/generate_imrd_docx_template.py \
    --title "中医药干预 COVID-19 恢复期症状改善的证据综合研究" \
    --author "张三" \
    --affiliation "某某大学中医学院" \
    --keywords "中医药; COVID-19; 证据矩阵; Gap Analysis" \
    --output-file output/IMRD_template.docx
"""

from __future__ import annotations

import argparse
import json
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List

from src.llm.llm_engine import LLMEngine

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError as exc:
    raise ImportError(
        "未检测到 python-docx。请先安装: c:/Users/hgk/tcmautoresearch/venv310/Scripts/python.exe -m pip install python-docx"
    ) from exc


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="生成 IMRD 论文模板（.docx）")
    parser.add_argument("--title", default="论文题目（请替换）", help="论文标题")
    parser.add_argument("--author", default="作者姓名（请替换）", help="作者")
    parser.add_argument("--affiliation", default="单位（请替换）", help="单位")
    parser.add_argument("--corresponding-email", default="通讯作者邮箱（请替换）", help="通讯作者邮箱")
    parser.add_argument("--keywords", default="关键词1; 关键词2; 关键词3", help="关键词，建议用分号分隔")
    parser.add_argument("--journal", default="目标期刊（请替换）", help="目标期刊")
    parser.add_argument("--language", default="zh", choices=["zh", "en"], help="模板语言")
    parser.add_argument("--research-domain", default="中医临床研究", help="研究领域，用于草稿生成")
    parser.add_argument("--clinical-question", default="", help="临床问题，用于草稿生成")
    parser.add_argument("--context-json", action="append", dest="context_json_paths", help="附加上下文 JSON 文件路径，可重复传入")
    parser.add_argument("--generate-draft", action="store_true", help="启用 Qwen 本地推理生成各章节草稿")
    parser.add_argument("--max-tokens", type=int, default=1024, help="草稿生成最大 token")
    parser.add_argument("--n-ctx", type=int, default=4096, help="模型上下文窗口")
    parser.add_argument("--temperature", type=float, default=0.2, help="采样温度")
    parser.add_argument("--draft-markdown-file", default="", help="可选：导出章节草稿 markdown 文件")
    parser.add_argument("--generate-figures", action="store_true", help="启用科研图包生成并接入论文文档")
    parser.add_argument("--figure-input-json", default="", help="图包输入 JSON，默认自动推断")
    parser.add_argument("--figure-output-dir", default="output/figure_pack", help="图包输出目录")
    parser.add_argument("--figure-dpi", type=int, default=300, help="图包导出 DPI")
    parser.add_argument("--figure-format", default="png", choices=["png", "tif", "tiff", "pdf"], help="图包导出格式")
    parser.add_argument("--embed-figures", action=argparse.BooleanOptionalAction, default=True, help="是否把图片嵌入 docx")
    parser.add_argument("--output-file", default="output/IMRD_template.docx", help="输出 docx 文件")
    return parser.parse_args()


def _set_default_font(document: DocumentType, language: str) -> None:
    style = document.styles["Normal"]
    font = style.font  # type: ignore[attr-defined]
    font.size = Pt(12)
    if language == "zh":
        font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # type: ignore[union-attr]
    else:
        font.name = "Times New Roman"


def _add_title_block(document: DocumentType, args: argparse.Namespace) -> None:
    p = document.add_paragraph(args.title)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(16)

    for line in [args.author, args.affiliation, f"Corresponding author: {args.corresponding_email}", f"Target journal: {args.journal}"]:
        p_line = document.add_paragraph(line)
        p_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    date_line = document.add_paragraph(f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def _add_heading_with_hint(document: DocumentType, heading: str, hints: Iterable[str], level: int = 1) -> None:
    document.add_heading(heading, level=level)
    for hint in hints:
        document.add_paragraph(hint)


def _add_common_sections_zh(document: DocumentType, args: argparse.Namespace) -> None:
    document.add_heading("摘要", level=1)
    document.add_paragraph("【背景】请简述研究背景与临床问题。")
    document.add_paragraph("【目的】请明确研究目的与核心科学问题。")
    document.add_paragraph("【方法】请说明研究设计、数据来源、统计方法。")
    document.add_paragraph("【结果】请总结主要结果。")
    document.add_paragraph("【结论】请总结临床意义与后续工作。")
    document.add_paragraph(f"关键词：{args.keywords}")

    _add_heading_with_hint(
        document,
        "1 引言（Introduction）",
        [
            "1.1 研究背景与临床需求。",
            "1.2 现有证据与研究缺口（Gap）。",
            "1.3 本研究目标与假设。",
        ],
    )

    _add_heading_with_hint(
        document,
        "2 方法（Methods）",
        [
            "2.1 研究设计（如回顾性研究、随机对照试验、证据综合等）。",
            "2.2 数据来源与检索策略（PubMed/MEDLINE、arXiv 等）。",
            "2.3 纳排标准与变量定义。",
            "2.4 证据矩阵构建与 Gap Analysis 流程。",
            "2.5 统计分析与软件。",
            "2.6 伦理声明。",
        ],
    )

    _add_heading_with_hint(
        document,
        "3 结果（Results）",
        [
            "3.1 文献纳入流程与来源分布。",
            "3.2 证据矩阵覆盖情况与主要维度命中。",
            "3.3 主要结局结果与敏感性分析。",
            "3.4 图表与附表引用（Figure 1-7 / Table 1-3）。",
        ],
    )

    _add_heading_with_hint(
        document,
        "4 讨论（Discussion）",
        [
            "4.1 主要发现与临床解释。",
            "4.2 与既往研究比较。",
            "4.3 局限性与潜在偏倚。",
            "4.4 临床转化价值与未来研究方向。",
        ],
    )

    _add_heading_with_hint(
        document,
        "5 结论（Conclusion）",
        [
            "请用 1-2 段概括核心贡献与临床意义。",
        ],
    )

    _add_heading_with_hint(
        document,
        "补充部分",
        [
            "作者贡献（Author Contributions）：[待填写]",
            "基金资助（Funding）：[待填写]",
            "利益冲突（Conflicts of Interest）：[待填写]",
            "数据可得性（Data Availability）：[待填写]",
            "致谢（Acknowledgments）：[待填写]",
        ],
    )

    document.add_heading("参考文献", level=1)
    document.add_paragraph("[1] 参考文献条目示例（按目标期刊格式补全）。")


def _add_common_sections_en(document: DocumentType, args: argparse.Namespace) -> None:
    _add_heading_with_hint(
        document,
        "Abstract",
        [
            "Background: Briefly describe the clinical context and unmet need.",
            "Objective: State the primary objective and hypothesis.",
            "Methods: Summarize design, data sources, and primary outcomes.",
            "Results: Report key findings with major statistics.",
            "Conclusion: Provide concise interpretation and implications.",
        ],
    )
    document.add_paragraph(f"Keywords: {args.keywords}")

    _add_heading_with_hint(
        document,
        "1. Introduction",
        [
            "1.1 Clinical background and rationale.",
            "1.2 Prior evidence and research gaps.",
            "1.3 Research objectives and hypothesis statements.",
        ],
    )

    _add_heading_with_hint(
        document,
        "2. Methods",
        [
            "2.1 Study design (e.g., retrospective, RCT, evidence synthesis).",
            "2.2 Data source and retrieval strategy (PubMed/MEDLINE, etc.).",
            "2.3 Eligibility criteria and variable definitions.",
            "2.4 Evidence matrix construction and gap-analysis prompt engineering.",
            "2.5 Statistical analysis and software.",
            "2.6 Ethics statement.",
        ],
    )

    _add_heading_with_hint(
        document,
        "3. Results",
        [
            "3.1 Literature screening and inclusion flow.",
            "3.2 Evidence matrix coverage and source comparison.",
            "3.3 Main outcomes and secondary analyses.",
            "3.4 Figure/Table callouts (Figure 1-7).",
        ],
    )

    _add_heading_with_hint(
        document,
        "4. Discussion",
        [
            "4.1 Principal findings and clinical interpretation.",
            "4.2 Comparison with prior studies.",
            "4.3 Strengths, limitations, and bias considerations.",
            "4.4 Future directions and translational potential.",
        ],
    )

    _add_heading_with_hint(
        document,
        "5. Conclusion",
        [
            "Concise summary of contribution and actionable implications.",
        ],
    )

    _add_heading_with_hint(
        document,
        "Supplementary Sections",
        [
            "Author Contributions: [to be filled]",
            "Funding: [to be filled]",
            "Conflicts of Interest: [to be filled]",
            "Data Availability: [to be filled]",
            "Acknowledgments: [to be filled]",
        ],
    )

    document.add_heading("References", level=1)
    document.add_paragraph("[1] Example reference item (format to target journal style).")


def _add_draft_sections_zh(document: DocumentType, args: argparse.Namespace, drafts: Dict[str, str]) -> None:
    document.add_heading("摘要", level=1)
    document.add_paragraph(drafts.get("Abstract", ""))
    document.add_paragraph(f"关键词：{args.keywords}")

    document.add_heading("1 引言（Introduction）", level=1)
    document.add_paragraph(drafts.get("Introduction", ""))

    document.add_heading("2 方法（Methods）", level=1)
    document.add_paragraph(drafts.get("Methods", ""))

    document.add_heading("3 结果（Results）", level=1)
    document.add_paragraph(drafts.get("Results", ""))

    document.add_heading("4 讨论（Discussion）", level=1)
    document.add_paragraph(drafts.get("Discussion", ""))

    document.add_heading("5 结论（Conclusion）", level=1)
    document.add_paragraph(drafts.get("Conclusion", ""))

    _add_heading_with_hint(
        document,
        "补充部分",
        [
            "作者贡献（Author Contributions）：[待填写]",
            "基金资助（Funding）：[待填写]",
            "利益冲突（Conflicts of Interest）：[待填写]",
            "数据可得性（Data Availability）：[待填写]",
            "致谢（Acknowledgments）：[待填写]",
        ],
    )

    document.add_heading("参考文献", level=1)
    document.add_paragraph("[1] 参考文献条目示例（按目标期刊格式补全）。")


def _add_draft_sections_en(document: DocumentType, args: argparse.Namespace, drafts: Dict[str, str]) -> None:
    document.add_heading("Abstract", level=1)
    document.add_paragraph(drafts.get("Abstract", ""))
    document.add_paragraph(f"Keywords: {args.keywords}")

    document.add_heading("1. Introduction", level=1)
    document.add_paragraph(drafts.get("Introduction", ""))

    document.add_heading("2. Methods", level=1)
    document.add_paragraph(drafts.get("Methods", ""))

    document.add_heading("3. Results", level=1)
    document.add_paragraph(drafts.get("Results", ""))

    document.add_heading("4. Discussion", level=1)
    document.add_paragraph(drafts.get("Discussion", ""))

    document.add_heading("5. Conclusion", level=1)
    document.add_paragraph(drafts.get("Conclusion", ""))

    _add_heading_with_hint(
        document,
        "Supplementary Sections",
        [
            "Author Contributions: [to be filled]",
            "Funding: [to be filled]",
            "Conflicts of Interest: [to be filled]",
            "Data Availability: [to be filled]",
            "Acknowledgments: [to be filled]",
        ],
    )

    document.add_heading("References", level=1)
    document.add_paragraph("[1] Example reference item (format to target journal style).")


def _load_context_payload(args: argparse.Namespace) -> Dict[str, Any]:
    payload: Dict[str, Any] = {
        "title": args.title,
        "author": args.author,
        "affiliation": args.affiliation,
        "journal": args.journal,
        "research_domain": args.research_domain,
        "clinical_question": args.clinical_question,
        "keywords": args.keywords,
        "language": args.language,
        "context_files": {},
    }

    for path_str in args.context_json_paths or []:
        path = Path(path_str)
        if not path.exists():
            payload["context_files"][path_str] = {"error": "file_not_found"}
            continue
        try:
            payload["context_files"][path_str] = json.loads(path.read_text(encoding="utf-8"))
        except Exception as exc:
            payload["context_files"][path_str] = {"error": f"json_parse_failed: {exc}"}

    return payload


def _generate_imrd_drafts(args: argparse.Namespace, context_payload: Dict[str, Any]) -> Dict[str, str]:
    section_prompts: List[str] = [
        "Abstract",
        "Introduction",
        "Methods",
        "Results",
        "Discussion",
        "Conclusion",
    ]

    llm = LLMEngine(
        max_tokens=args.max_tokens,
        n_ctx=args.n_ctx,
        temperature=args.temperature,
        verbose=False,
    )
    llm.load()
    try:
        drafts: Dict[str, str] = {}
        for section_name in section_prompts:
            section_context = {
                **context_payload,
                "section_requirements": {
                    "Abstract": "结构化摘要，包含背景/目的/方法/结果/结论。",
                    "Introduction": "突出临床需求、证据缺口和研究问题。",
                    "Methods": "说明检索策略、证据矩阵、统计方案与伦理。",
                    "Results": "报告主要发现，注意与图表编号衔接。",
                    "Discussion": "解释意义、比较文献、局限性与未来方向。",
                    "Conclusion": "给出简洁可执行结论。",
                }.get(section_name, ""),
            }
            drafts[section_name] = llm.draft_section(section_name, section_context)
        return drafts
    finally:
        llm.unload()


def _export_drafts_markdown(args: argparse.Namespace, drafts: Dict[str, str]) -> str:
    output_path = Path(args.draft_markdown_file) if args.draft_markdown_file else Path(args.output_file).with_suffix(".md")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    lines = ["# IMRD 初稿（Qwen 本地推理）", ""]
    for section in ["Abstract", "Introduction", "Methods", "Results", "Discussion", "Conclusion"]:
        lines.append(f"## {section}")
        lines.append(drafts.get(section, ""))
        lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")
    return output_path.as_posix()


def _resolve_figure_input_json(args: argparse.Namespace) -> Path:
    if args.figure_input_json:
        return Path(args.figure_input_json)

    for path_str in args.context_json_paths or []:
        candidate = Path(path_str)
        if candidate.exists():
            return candidate

    return Path("output/local_tcm_llm_analysis.json")


def _run_figure_generation(args: argparse.Namespace) -> Dict[str, Any]:
    figure_script = Path(__file__).resolve().parent / "generate_submission_figure_pack.py"
    input_json = _resolve_figure_input_json(args)
    output_dir = Path(args.figure_output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    command = [
        sys.executable,
        figure_script.as_posix(),
        "--input-json",
        input_json.as_posix(),
        "--output-dir",
        output_dir.as_posix(),
        "--dpi",
        str(args.figure_dpi),
        "--format",
        args.figure_format,
    ]

    completed = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", errors="replace")
    if completed.returncode != 0:
        stderr_text = (completed.stderr or "").strip()
        raise RuntimeError(f"科研图包生成失败: {stderr_text[-800:]}")

    parsed_output: Dict[str, Any] = {}
    stdout_text = (completed.stdout or "").strip()
    if stdout_text:
        try:
            parsed_output = json.loads(stdout_text)
        except json.JSONDecodeError:
            parsed_output = {}

    figure_files = sorted(str(path.as_posix()) for path in output_dir.glob(f"Figure*.{args.figure_format}"))
    if not parsed_output:
        parsed_output = {
            "input_json": input_json.as_posix(),
            "output_dir": output_dir.as_posix(),
            "dpi": args.figure_dpi,
            "format": args.figure_format,
            "figure_pack": figure_files,
            "legends_file": (output_dir / "Figure_legends_submission.md").as_posix(),
        }
    elif not parsed_output.get("figure_pack"):
        parsed_output["figure_pack"] = figure_files

    return parsed_output


def _add_figure_section(document: DocumentType, args: argparse.Namespace, figure_result: Dict[str, Any]) -> None:
    is_zh = args.language == "zh"
    document.add_heading("图表与图注" if is_zh else "Figures and Legends", level=1)

    output_dir = figure_result.get("output_dir", args.figure_output_dir)
    legends_file = figure_result.get("legends_file", "")
    figure_pack = figure_result.get("figure_pack", []) or []

    if is_zh:
        document.add_paragraph(f"图包目录：{output_dir}")
        if legends_file:
            document.add_paragraph(f"图注文件：{legends_file}")
    else:
        document.add_paragraph(f"Figure directory: {output_dir}")
        if legends_file:
            document.add_paragraph(f"Legend file: {legends_file}")

    if not figure_pack:
        document.add_paragraph("未发现图包文件。" if is_zh else "No figure files were generated.")
        return

    embeddable_suffixes = {".png", ".jpg", ".jpeg", ".bmp", ".gif"}
    for idx, figure_path in enumerate(figure_pack, start=1):
        path_obj = Path(figure_path)
        document.add_paragraph(f"Figure {idx}: {path_obj.name}")
        if args.embed_figures and path_obj.suffix.lower() in embeddable_suffixes and path_obj.exists():
            document.add_picture(path_obj.as_posix())
        elif args.embed_figures and path_obj.suffix.lower() not in embeddable_suffixes:
            if is_zh:
                document.add_paragraph("该格式不支持直接嵌入 docx，请在投稿阶段手动插入。")
            else:
                document.add_paragraph("This format is not embeddable in docx directly. Please insert manually for submission.")


def main() -> None:
    args = parse_args()
    document = Document()
    _set_default_font(document, args.language)
    _add_title_block(document, args)

    draft_markdown = ""
    if args.generate_draft:
        context_payload = _load_context_payload(args)
        drafts = _generate_imrd_drafts(args, context_payload)
        if args.language == "zh":
            _add_draft_sections_zh(document, args, drafts)
        else:
            _add_draft_sections_en(document, args, drafts)
        draft_markdown = _export_drafts_markdown(args, drafts)
    else:
        if args.language == "zh":
            _add_common_sections_zh(document, args)
        else:
            _add_common_sections_en(document, args)

    figure_result: Dict[str, Any] = {}
    if args.generate_figures:
        figure_result = _run_figure_generation(args)
        _add_figure_section(document, args, figure_result)

    output_path = Path(args.output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path.as_posix())
    if args.generate_draft:
        print(f"IMRD 初稿已生成: {output_path.as_posix()}")
        print(f"章节草稿 Markdown: {draft_markdown}")
    else:
        print(f"IMRD 模板已生成: {output_path.as_posix()}")

    if args.generate_figures:
        print(f"科研图包目录: {figure_result.get('output_dir', args.figure_output_dir)}")
        print(f"科研图数量: {len(figure_result.get('figure_pack', []) or [])}")
        legends_file = figure_result.get("legends_file", "")
        if legends_file:
            print(f"图注文件: {legends_file}")


if __name__ == "__main__":
    main()
